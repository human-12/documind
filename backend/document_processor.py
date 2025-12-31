from typing import List, Dict
import os
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl
from langchain.text_splitter import RecursiveCharacterTextSplitter
import hashlib

class DocumentProcessor:
    """Process documents and extract text content"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def extract_text(self, file_path: str, file_type: str) -> tuple[str, Dict]:
        """Extract text from different file types"""
        if file_type == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_type == "docx":
            return self._extract_from_docx(file_path)
        elif file_type == "xlsx":
            return self._extract_from_xlsx(file_path)
        elif file_type == "txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> tuple[str, Dict]:
        """Extract text from PDF"""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
        
        metadata = {
            "page_count": len(reader.pages),
            "file_type": "pdf"
        }
        return text.strip(), metadata
    
    def _extract_from_docx(self, file_path: str) -> tuple[str, Dict]:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "file_type": "docx"
        }
        return text.strip(), metadata
    
    def _extract_from_xlsx(self, file_path: str) -> tuple[str, Dict]:
        """Extract text from XLSX"""
        workbook = openpyxl.load_workbook(file_path)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n\n=== Sheet: {sheet_name} ===\n\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"
        
        metadata = {
            "sheet_count": len(workbook.sheetnames),
            "file_type": "xlsx"
        }
        return text.strip(), metadata
    
    def _extract_from_txt(self, file_path: str) -> tuple[str, Dict]:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        metadata = {
            "file_type": "txt"
        }
        return text.strip(), metadata
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)
        
        chunk_documents = []
        for i, chunk in enumerate(chunks):
            chunk_doc = {
                "content": chunk,
                "chunk_index": i,
                "metadata": metadata or {}
            }
            chunk_documents.append(chunk_doc)
        
        return chunk_documents
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Get file type from filename"""
        extension = Path(filename).suffix.lower()
        type_mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".xlsx": "xlsx",
            ".xls": "xlsx",
            ".txt": "txt"
        }
        return type_mapping.get(extension, "unknown")
    
    @staticmethod
    def calculate_file_hash(file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
